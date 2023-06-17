from django.shortcuts import render, HttpResponse, HttpResponseRedirect
from django.http import JsonResponse
import random, re, datetime
import docx
from docx.shared import Pt
from .models import Task_shablon, Obj_shablon, Tasks

# Create your views here.
def index(request):
    return render(request, 'main/index.html')

def about(request):
    return render(request, 'main/about.html')

def reg(request):
    return render(request, 'main/reg.html')

def generate_tasks(request):
    return render(request, 'main/create_tasks.html')

def generated_tasks(request):
    # if request.method == "POST":
    #     theme_tasks = request.POST.get("theme_tasks")
    #     return JsonResponse(theme_tasks)
    # else:
    #     # Обычный HTTP-запрос
    #     data = {'name': 'tA', 'phone': 'BLYAT'}
    #     return HttpResponse('test.html')
    return render(request, 'main/generated_tasks.html')

def kid_generate_tasks(request):
    return render(request, 'main/create_tasks_kid.html')


def test(request):
    ##insert_shablons_electro(obj)
    response = HttpResponse()
    url = "/static/main/img/children.png"
    response = url
    return response
    #return render(request, 'main/test2.html', {"task_shablons": rdy_task})

def upload(request):
    return render(request, 'main/upload.html')

def tested_ajax(request):
    # получаем из данных запроса POST отправленные через форму данные
    theme = request.POST.get("choose_theme_dropbox", "Undefined")
    return HttpResponse(f"<h2>В DB было : {theme}</h2>")

def is_ajax(request):
    return request.META.get('HTTP_X_REQUESTED_WITH') == 'XMLHttpRequest'

def my_view(request):
    if request.method == "GET":
        # data = {'name': request.name, 'phone': 2292747}
        theme_tasks = {'name':request.GET.get("name")}
        return JsonResponse(theme_tasks)
    else:
        # Обычный HTTP-запрос
        return HttpResponse('test.html')

def add(request):
    if request.method == "POST":
        if request.POST.get("v_fio_input").isdigit():
            count_vars = int(request.POST.get("v_fio_input"))
        else:
            count_vars = 1
        all_tasks = {}
        if request.POST.get("kin_input").isdigit():
            all_tasks["kinematic"] = int(request.POST.get("kin_input"))
        if request.POST.get("elec_input").isdigit():
            all_tasks["electro"] = int(request.POST.get("elec_input"))
        if request.POST.get("press_input").isdigit():
            all_tasks["pressure"] = int(request.POST.get("press_input"))
        if request.POST.get("pow_input").isdigit():
            all_tasks["power"] = int(request.POST.get("pow_input"))
        res = varaints(count_vars,all_tasks)
        return HttpResponse('upload.html')
        #return render(request, 'main/test2.html', {"task_shablons": res})
        #return JsonResponse({1:1})
    return HttpResponseRedirect("test.html")

#variants
def varaints(count_variants,theme_task_dict):#количество вариантов и словарь тема:количество заданий
    result_list = []
    for i in range(count_variants):
        result_list += gen_tasks(theme_task_dict)
    # result_list
    return result_list


def gen_tasks(theme_task_dict):
    result_list = []
    for k,v in theme_task_dict.items():
        if k == "kinematic":
            result_list += generate_kinematic(v)
        elif k == "electro":
            result_list += generate_kinematic(v)
        elif k == "power":
            result_list += generate_power(v)
        elif k == "pressure":
            result_list += generate_pressure(v)
    result_dict = dict(zip(result_list, list(range(1,len(result_list)+1))))
    return result_dict


def generate_pressure(count_tasks):
    result = []
    for i in range(count_tasks):
        task_shablons = Task_shablon.objects.all()
        task_on_theme = task_shablons.filter(theme = "pressure")
        tasks_list = task_on_theme.values_list()
        task = tasks_list[random.randint(0,len(task_on_theme)-1)][2]
        obj = insert_shablons_pressure(task)
        result.append(obj)
    return result


def insert_shablons_pressure(task):
    objects = {'бензина' : 710, 'нефти': 800, 'воды': 1000, 'ртути' : 13600,'керосина': 800, 'эфира': 710} # плотности веществ
    using_object = list(objects.keys())[random.randint(0, len(objects) - 1)] #выбираем случайный объект
    object_value = objects[using_object]
    height_for_task = round(random.randint(1,50) * 0.1, 1)
    pressure = random.randint(2,35)
    result = task
    result = re.sub(r'OBJECT', using_object, result)
    result = re.sub(r'DENSITY',str(object_value )+' кг/м^3',result)
    result = re.sub(r'HEIGHT', str(height_for_task) + ' м', result)
    result = re.sub(r'PRESSURE', str(pressure) + ' кПа', result)
    return result


#power generate
def generate_power(count_tasks):
    result = []
    for i in range(count_tasks):
        task_shablons = Task_shablon.objects.all()
        task_on_theme = task_shablons.filter(theme = "power")
        tasks_list = task_on_theme.values_list()
        task = tasks_list[random.randint(0,len(task_on_theme)-1)][2]
        obj = insert_shablons_power(task)
        result.append(obj)
    return result

def insert_shablons_power(task):
    objects = {'мотороллера' : [[35, 90], [150,300], [1450,7500]], 'грузовика': [[75, 120], [550,1200],[41000, 144000]], 'самолета' : [[600, 900],[600000,775000],[3600000, 6975000]]} #[[range(velocity), range(power)]]
    using_object = list(objects.keys())[random.randint(0, len(objects) - 1)]#выбираем случайный объект
    object_value_range = objects[using_object]
    object_value_velocity = random.randint(object_value_range[0][0], object_value_range[0][1])
    object_value_force = random.randint(object_value_range[1][0], object_value_range[1][1])
    object_value_power = random.randint(object_value_range[2][0], object_value_range[2][1])
    result = task
    result = re.sub(r'OBJECT', using_object, result)
    result = re.sub(r'VELOCITY',str(object_value_velocity )+' км / ч',result)
    result = re.sub(r'FORCE', str(object_value_force) + ' Н', result)
    result = re.sub(r'POWER', str(object_value_power) + ' Вт', result)
    # result = re.sub(r'VOLTAGE', str(voltage) + ' В', result)
    return result

# electro generate
def generate_electro(count_tasks):
    result = []
    for i in range(count_tasks):
        task_shablons = Task_shablon.objects.all()
        task_on_theme = task_shablons.filter(theme = "electro")
        tasks_list = task_on_theme.values_list()
        task = tasks_list[random.randint(0,len(task_on_theme)-1)][2]
        obj = insert_shablons_electro(task)
        result.append(obj)
    return result


def insert_shablons_electro(task):
    amper = {'мА' : [10, 100], 'А': [1, 300]}
    amper_unit = list(amper.keys())[random.randint(0, len(amper) - 1)]#выбираем единицу измерения из А и мА
    amper_value_range = amper[amper_unit]
    amper_value = round(random.randint(amper_value_range[0], amper_value_range[1]) * 0.1, 1)
    voltage = round(random.randint(5,2200) * 0.1,1)
    resistance = round(random.randint(1,4800) * 0.1,1)
    result = task
    result = re.sub(r'RESISTANCE',str(resistance)+' Ом',result)
    result = re.sub(r'AMPER', str(amper_value) + ' ' + amper_unit, result)
    result = re.sub(r'VOLTAGE', str(voltage) + ' В', result)
    return result

#kinematic
def insert_shablons_kinematic(task):
    obj_shablon = Obj_shablon.objects.all()
    obj_list = obj_shablon.values_list()
    obj = obj_list[random.randint(0, len(obj_list) - 1)]
    time = obj[5].split('/')
    if time[1] == 'С‡': #Если часы, просто там свои приколы с кодировкой, с которой очень лень разбираться
        time_value = random.randint(1, 20)
    else:
        time_value = random.randint(1, 60)
    distance = random.randint(5, int(obj[6]))
    speed = random.randint(int(obj[3]), int(obj[4]))
    result = re.sub(r'OBJ',obj[2],task)
    result = re.sub(r'T',str(time_value)+' '+ time[1],result)
    result = re.sub(r'V', str(speed) + ' ' + time[0] + ' / ' + time[1], result)
    result = re.sub(r'S', str(distance) + ' ' + time[0], result)
    return result


# def save_in_doc(list_tasks, var):
#     doc = docx.Document()
#     style = doc.styles['Normal']
#     style.font.name = 'Times New Roman'
#     style.font.size = Pt(14)
#     doc.add_paragraph('Вариант '+ str(var))
#     par1 = doc.add_paragraph()
#     for i in list_tasks:
#         par1.add_run(i + '\n')
#     name_file = str(datetime.datetime.now()).split('.')[0].split(':')
#     name_file = name_file[0]+'-'+name_file[1]+'-'+name_file[2]
#     doc.save("main/static/docs/"+name_file+'.docx')
#     task_in_doc = Tasks()
#     task_in_doc.document("main/static/docs/"+name_file+'.docx')
#     all_tasks = Tasks.objects.all()
#     tasks_list = all_tasks.values_list()
#     return tasks_list[-1]



def generate_kinematic(count_tasks):
    result = []
    for i in range(count_tasks):
        task_shablons = Task_shablon.objects.all()
        task_on_theme = task_shablons.filter(theme = "kinematica")
        tasks_list = task_on_theme.values_list()
        task = tasks_list[random.randint(0,len(task_on_theme)-1)][2]
        obj = insert_shablons_kinematic(task)
        result.append(obj)
    return result


def log_in(request):
    if request.method == "POST":
        if request.POST.get("input_login") == "s4ika":
            return HttpResponseRedirect('main/create_tasks.html')
        else:
            return HttpResponseRedirect( 'main/create_tasks_kid.html')


delitel = lambda value1, value2: value1 / value2
umnozhitel = lambda value1, value2: value1 * value2

def div_for_v(unit1,value1,unit2,value2):
    if (unit1 == "м" or unit1 == "метры"):
        return value1/value2
    else:
        return value2 / value1


def formulas_for_finding(value1,unit1,value2,unit2):
    #Уточнить в порядке ли будут проверяться или нет
    print(value1,unit1,value2,unit2)
    formls = {
        delitel(value1,value2): [['метры', 'с'], ['метры', 'м/c'],['В','А'], ['В','Ом'],['Вт','м/с'],['Вт','Н'],['Па','кг/м^3'],['Па','м']],
        umnozhitel(value1, value2): [['м/с', 'с'], ['Ом','А'], ['Н','м/с'],['кг/м^3','м']],
    }
    res = 0
    check_units = [unit1, unit2]
    print(check_units)
    for k, v in formls.items():
        if check_units in v:
            res = round(k,3)
    return res

#Приводим км,см к единицам СИ.Аналогичную бахну для времени и скорости
def si_units_s(value_unit):
    value = float(value_unit[0])
    units = {
        'сантиметров': value * 0.01,
        'см': value * 0.01,
        'километров': value * 1000,
        'км': value * 1000,
        'мм': value * 0.001
    }

    if value_unit[1] in units:
        value = units[value_unit[1]]
        unit = 'метры'
    else:
        unit = 'метры'
    print('В единицах СИ : ', value,' ',unit)
    return value,unit


def si_units_t(value_unit):
    value = float(value_unit[0])
    units = {
        'минут': value * 60,
        'минуты': value * 60,
        'мин': value * 60,
        'часов': value * 3600,
        'часа': value * 3600,
        'ч': value * 3600,
    }

    if value_unit[1] in units:
        value = units[value_unit[1]]
        unit = 'с'
    else:
        unit = 'с'
    print('В единицах СИ : ', value,' ',unit)
    return value,unit


def si_units_v(value_unit):
    print(value_unit)
    value = float(value_unit[0])
    units = {
        'км / ч': value * 0.2777,
        'километров в час': value * 0.2777,
        'километров в минуту': value * 16.66,
        'км / мин': value * 16.66,
        'метров в минуту': value / 60,
        'м / мин': value * 3600,
        'сантиметров в секунду': value / 100,
        'см / с': value / 100
    }
    if value_unit[1] in units:
        value = units[value_unit[1]]
        unit = 'м/с'
    else:
        unit = 'м/с'
    print('В единицах СИ : ', value,' ',unit)
    return value,unit


#С помощью регулярки ищем данные
def find_value_metr_and_time(s):
    si_value_metr = unit_metr = si_value_v = si_unit_v = si_value_time = si_unit_time = 0
    velocity = re.findall(r'\d+\s[метровкилсан]+\s[в/]\s[часекундмит]+', s) #Сначала лучше проверить, есть ли значение скорости в задаче,
    # т.к. у нее в ЕИ используются единицы длинны и времени одновременно
    if len(velocity) > 0: #Проверяем нашли ли информацию о скорости объекта
        s = re.sub(r'\d+\s[метровкилсан]+\s[в/]\s[часекундмит]+', '',s)  # Вырезаем velocity из задачи, чтобы лишний раз его не трогать
        temp = velocity[0]
        value_unit = re.split(r' ', temp) #Разбили по пробелам
        unit = ''
        for i in value_unit[1:len(value_unit) + 1]: #Цикл для того чтобы объеденить строки типа ['км'] ['в'] ['час']
            unit += i + ' '
        unit = unit[:len(unit) - 1] #Удаляем лишний пробел
        velocity_value_unit = [value_unit[0], unit]
        si_value_v,si_unit_v = si_units_v(velocity_value_unit)
    print("Строка после VELOCITY = ",s)
    spacing = re.findall(r'\d+\s[санти]*[кило]*[с]?[к]?м[етров]*', s)
    if len(spacing) > 0:
        s = re.sub(r'\d+\s[санти]*[кило]*[с]?[к]?м[етров]*', '', s) # Избавляемся от него в тексте задачи
        temp = spacing[0]
        value_unit = re.split(r' ', temp)
        si_value_metr,unit_metr = si_units_s(value_unit)
    time = re.findall(r'\d+\s[секундаы]*[минутаы]*[час]*', s)#ищем время и его значение
    if len(time) > 0:
        temp = time[0]
        value_unit = re.split(r' ', temp)#разбиваем полученную строку
    #print(value_unit)
        si_value_time, si_unit_time =si_units_t(value_unit)
    if len(velocity) > 0:
        if len(spacing) > 0:
            return formulas_for_finding(si_value_metr,unit_metr,si_value_v,si_unit_v)
        else:
            return formulas_for_finding(si_value_v, si_unit_v, si_value_time, si_unit_time)
    else:
        return formulas_for_finding(si_value_metr, unit_metr, si_value_time, si_unit_time)


def find_value_time(s):
    result = re.findall(r'\d+\s[секундаы]*[минутаы]*[час]*', s)
    print(result)
    temp = result[0]
    value_unit = re.split(r' ', temp)
    print(value_unit)


def si_units_amper(value_unit):
    value = float(value_unit[0])
    units = {
        'А': value,
        'мА': value * 0.001
    }
    if value_unit[1] in units:
        value = units[value_unit[1]]
        unit = 'А'
    else:
        unit = 'А'
    print('В единицах СИ : ', value,' ',unit)
    return value,unit



def electro_solutions(s):
    si_value_amper = unit_amper = si_value_voltage = si_unit_voltage = si_value_resistance = si_unit_resistance = 0
    amper = re.findall(r'\d+[.\d]*\s[Аампер]+',s)  # Сначала лучше проверить, есть ли значение amper в задаче,
    if len(amper) > 0:  # Проверяем нашли ли информацию о скорости объекта
        s = re.sub(r'\d+[.\d]*\s[Аампер]+','',s)  # Вырезаем amper из задачи, чтобы лишний раз его не трогать
        temp = amper[0]
        value_unit = re.split(r' ', temp)  # Разбили по пробелам
        unit = ''
        for i in value_unit[1:len(value_unit) + 1]:  # Цикл для того чтобы объеденить строки типа ['км'] ['в'] ['час']
            unit += i + ' '
        unit = unit[:len(unit) - 1]  # Удаляем лишний пробел
        amper_value_unit = [value_unit[0], unit]
        si_value_amper, si_unit_amper = si_units_amper(amper_value_unit)
    voltage = re.findall(r'\d+[.\d]*\s[кВвольт]+',s)
    if len(voltage) > 0:
        s = re.sub(r'\d+[.\d]*\s[кВвольт]+', '', s)  # Избавляемся от него в тексте задачи
        temp = voltage[0]
        value_unit = re.split(r' ', temp)
        si_value_voltage = float(value_unit[0])
        si_unit_voltage = value_unit[1]
        # si_value_metr, unit_metr = si_units_s(value_unit)
    resistance = re.findall(r'\d+[.\d]*\s[Ом]+', s)  # ищем время и его значение
    if len(resistance) > 0:
        temp = resistance[0]
        value_unit = re.split(r' ', temp)  # разбиваем полученную строку
        # print(value_unit)
        si_value_resistance  = float(value_unit[0])
        si_unit_resistance = value_unit[1]
    if len(amper) > 0:
        if len(voltage) > 0:
            return formulas_for_finding(si_value_voltage, si_unit_voltage, si_value_amper, si_unit_amper)
        else:
            return formulas_for_finding(si_value_resistance, si_unit_resistance, si_value_amper, si_unit_amper)
    else:
        return formulas_for_finding(si_value_voltage, si_unit_voltage, si_value_resistance, si_unit_resistance)


def power_solutions(s):
    si_value_force = si_unit_force = si_value_velocity = si_unit_velocity = si_value_power = si_unit_power = 0
    velocity = re.findall(r'\d+\s[метровкилсан]+\s[в/]\s[часекундмит]+', s)  # Сначала лучше проверить, есть ли значение amper в задаче,
    print(velocity)
    if len(velocity) > 0:  # Проверяем нашли ли информацию о скорости объекта
        s = re.sub(r'\d+\s[метровкилсан]+\s[в/]\s[часекундмит]+','',s)  # Вырезаем amper из задачи, чтобы лишний раз его не трогать
        print(s)
        temp = velocity[0]
        value_unit = re.split(r' ', temp)  # Разбили по пробелам
        unit = ''
        for i in value_unit[1:len(value_unit) + 1]:  # Цикл для того чтобы объеденить строки типа ['км'] ['в'] ['час']
            unit += i + ' '
        unit = unit[:len(unit) - 1]  # Удаляем лишний пробел
        velocity_value_unit = [value_unit[0], unit]
        si_value_velocity, si_unit_velocity = si_units_v(velocity_value_unit)
    force = re.findall(r'\d+[.\d]*\s[кмГМНьютон]+',s)
    print(force)
    if len(force) > 0:
        s = re.sub(r'\d+[.\d]*\s[кмГМНьютон]+', '', s)  # Избавляемся от него в тексте задачи
        temp = force[0]
        value_unit = re.split(r' ', temp)
        si_value_force = float(value_unit[0])
        si_unit_force = value_unit[1]
    power = re.findall(r'\d+[.\d]*\s[кМГВат]+',s)  # ищем время и его значение
    if len(power) > 0:
        temp = power[0]
        value_unit = re.split(r' ', temp)  # разбиваем полученную строку
        # print(value_unit)
        si_value_power  = float(value_unit[0])
        si_unit_power = value_unit[1]
    if len(velocity) > 0:
        if len(force) > 0:
            return formulas_for_finding(si_value_force, si_unit_force, si_value_velocity, si_unit_velocity)
        else:
            return formulas_for_finding(si_value_power, si_unit_power, si_value_velocity, si_unit_velocity)
    else:
        return formulas_for_finding(si_value_power, si_unit_power, si_value_force, si_unit_force)


def si_units_density(value_unit):
    print(value_unit)
    value = float(value_unit[0])
    units = {
        'кг/м^3': value,
        'г/см^3': value * 1000,
        'г/м^3': value / 1000,
    }
    if value_unit[1] in units:
        value = units[value_unit[1]]
        unit = 'кг/м^3'
    else:
        unit = 'кг/м^3'
    print('В единицах СИ : ', value,' ',unit)
    return value,unit


def si_units_pressure(value_unit):
    print(value_unit)
    value = float(value_unit[0])
    units = {
        'кПа': value * 1000,
        'ГПа': value * 1000000000,
        'МПа': value / 1000000,
    }
    if value_unit[1] in units:
        value = units[value_unit[1]]
        unit = 'Па'
    else:
        unit = 'Па'
    print('В единицах СИ : ', value,' ',unit)
    return value,unit


def pressure_solutions(s):
    si_value_pressure = si_unit_pressure = si_value_density = si_unit_density = si_value_height = si_unit_height = 0
    density = re.findall(r'\d+[.\d]*\s[килограмтн]+[в/]+[смк^3]+', s)  # Сначала лучше проверить, есть ли значение amper в задаче,
    print(density)
    if len(density) > 0:  # Проверяем нашли ли информацию о скорости объекта
        s = re.sub(r'\d+[.\d]*\s[килограмтн]+[в/]+[смк^3]+','',s)  # Вырезаем amper из задачи, чтобы лишний раз его не трогать
        print(s)
        temp = density[0]
        value_unit = re.split(r' ', temp)  # Разбили по пробелам
        unit = ''
        for i in value_unit[1:len(value_unit) + 1]:  # Цикл для того чтобы объеденить строки типа ['км'] ['в'] ['час']
            unit += i + ' '
        unit = unit[:len(unit) - 1]  # Удаляем лишний пробел
        density_value_unit = [value_unit[0], unit]
        si_value_density, si_unit_density = si_units_density(density_value_unit)
    height = re.findall(r'\d+[.\d]*\s[санти]*[кило]*[с]?[к]?м[етров]*', s)
    print(height)
    if len(height) > 0:
        s = re.sub(r'\d+[.\d]*\s[санти]*[кило]*[с]?[к]?м[етров]*', '', s)  # Избавляемся от него в тексте задачи
        temp = height[0]
        value_unit = re.split(r' ', temp)
        si_value_height = float(value_unit[0])
        si_unit_height = value_unit[1]
    pressure = re.findall(r'\d+[.\d]*\s[ГМПаскль]+', s)
    if len(pressure) > 0:
        temp = pressure[0]
        value_unit = re.split(r' ', temp)  # разбиваем полученную строку
        # print(value_unit)
        si_value_pressure, si_unit_pressure = si_units_pressure(value_unit)
    if len(density) > 0:
        if len(height) > 0:
            return formulas_for_finding(si_value_density * 10, si_unit_density, si_value_height, si_unit_height)
        else:
            return formulas_for_finding(si_value_pressure, si_unit_pressure, si_value_density * 10, si_unit_density)
    else:
        return formulas_for_finding(si_value_pressure, si_unit_pressure, si_value_height * 10, si_unit_height)

#Считываем текст из документа и возвращаем список задач
def read_docx_file(path):
    doc = docx.Document(path)
    prev_text = []
    for paragraph in doc.paragraphs:
        prev_text.append(paragraph.text)
    text = [prev_text[0]]
    temp = prev_text[1].split('\n')
    for i in temp:
        if i != '':
            text += [i[3:]]
    return text

#Вызываем для каждой темы задачи необходимые методы для ее решения
#На вход список задач и словарь {"тема": кол-во задач по теме,...}
def get_answrs_on_theme(list_tasks,theme):
    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    doc.add_paragraph('Ответы на ' + list_tasks[0])
    num = 1
    results_text = []
    for key in theme:
        if key == "Кинематика":
            for i in range(theme[key]):
                res = str(num) + '. ' + str(find_value_metr_and_time(list_tasks[num]))
                num += 1
                results_text.append(res)
        elif key == "Электродинамика":
            for i in range(theme[key]):
                res = str(num) + '. ' + str(electro_solutions(list_tasks[num]))
                num += 1
                results_text.append(res)
        elif key == "Мощность":
            for i in range(theme[key]):
                res = str(num) + '. ' + str(power_solutions(list_tasks[num]))
                num += 1
                results_text.append(res)
        elif key == "Давление":
            for i in range(theme[key]):
                res = str(num) + '. ' + str(pressure_solutions(list_tasks[num]))
                num += 1
                results_text.append(res)
    par1 = doc.add_paragraph()
    for i in results_text:
        par1.add_run(i + '\n')
    name_file =  str(datetime.datetime.now()).split('.')[0].split(':')
    name_file = 'Answers' + name_file[0] + '-' + name_file[1] + '-' + name_file[2]
    doc.save(name_file + '.docx')